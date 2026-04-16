import io
import logging
from typing import List, Optional

from .base import BaseFileProcessor, ProcessorResult

logger = logging.getLogger(__name__)


class DocxProcessor(BaseFileProcessor):
    async def run_baseline(
        self,
        file_bytes: bytes,
        media_file,
        tier: int,
        analysis_kinds: List[str],
        call_context: Optional[object] = None,
    ) -> List[ProcessorResult]:
        if tier != 0:
            return []

        from docx import Document
        from media_storage.constants import MAX_TEXT_BYTES

        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception as exc:
            logger.warning('Failed to open DOCX %s: %s', media_file.uuid, exc)
            return []

        results: list[ProcessorResult] = []
        if 'docx_text_extraction' in analysis_kinds:
            results.append(self._extract_text(document, MAX_TEXT_BYTES))
        if 'docx_structure' in analysis_kinds:
            results.append(self._extract_structure(document))
        return results

    def _extract_text(self, document, max_bytes: int) -> ProcessorResult:
        paragraphs: list[str] = []
        total_chars = 0

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            paragraphs.append(text)
            total_chars += len(text)
            if total_chars > max_bytes:
                break

        full_text = '\n\n'.join(paragraphs)
        truncated = total_chars > max_bytes
        return ProcessorResult(
            analysis_kind='docx_text_extraction',
            summary_text=f'Extracted {total_chars:,} chars from DOCX' + (' (truncated)' if truncated else ''),
            result_json={
                'text': full_text[:500_000],
                'total_chars': total_chars,
                'paragraph_count': len(paragraphs),
                'truncated': truncated,
            },
        )

    def _extract_structure(self, document) -> ProcessorResult:
        headings = []
        for paragraph in document.paragraphs:
            style_name = getattr(getattr(paragraph, 'style', None), 'name', '') or ''
            if style_name.startswith('Heading') and paragraph.text.strip():
                headings.append(
                    {
                        'level': style_name,
                        'text': paragraph.text[:100],
                    }
                )

        image_count = 0
        for rel in document.part.rels.values():
            if 'image' in (rel.reltype or ''):
                image_count += 1

        result = {
            'heading_count': len(headings),
            'headings': headings[:50],
            'table_count': len(document.tables),
            'image_count': image_count,
            'paragraph_count': len(document.paragraphs),
        }
        summary = (
            f'DOCX: {len(document.paragraphs)} paragraphs, '
            f'{len(headings)} headings, '
            f'{len(document.tables)} tables, '
            f'{image_count} images'
        )
        return ProcessorResult(
            analysis_kind='docx_structure',
            summary_text=summary,
            result_json=result,
        )
